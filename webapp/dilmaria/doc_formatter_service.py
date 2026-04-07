from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from webapp.dilmaria.ai_structurer import AIStructurerService
from webapp.dilmaria.doc_formatter_schema import (
    DocFormatterPayload,
    DocFormatterResult,
    StructuredBlock,
)
from webapp.dilmaria.exceptions import AgentExecutionError

PLACEHOLDER = "{{CONTEUDO}}"
STYLE_CANDIDATES = {
    "title": ["Heading 1", "Titulo 1", "Título 1", "Normal"],
    "subtitle": ["Heading 2", "Titulo 2", "Título 2", "Normal"],
    "paragraph": ["Normal"],
}


async def run_doc_formatter_agent(raw_payload: dict) -> DocFormatterResult:
    payload = DocFormatterPayload.model_validate(raw_payload)
    document = Document(BytesIO(payload.template_bytes))

    blocks = await AIStructurerService().structure_text(payload.text)
    if payload.mode == "replace_body":
        _replace_document_body(document, blocks)
    else:
        placeholder_paragraph = _find_placeholder_paragraph(document)
        if placeholder_paragraph is None:
            raise AgentExecutionError(f"Placeholder obrigatorio nao encontrado: {PLACEHOLDER}")
        _apply_blocks(document, placeholder_paragraph, blocks)
    _enable_field_updates_on_open(document)

    output = BytesIO()
    document.save(output)
    return DocFormatterResult(document_bytes=output.getvalue())


def _find_placeholder_paragraph(document: DocumentObject) -> Paragraph | None:
    for paragraph in _iter_paragraphs(document):
        if PLACEHOLDER in paragraph.text:
            return paragraph
    return None


def _iter_paragraphs(document: DocumentObject):
    for paragraph in document.paragraphs:
        yield paragraph

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def _apply_blocks(
    document: DocumentObject,
    placeholder_paragraph: Paragraph,
    blocks: list[StructuredBlock],
) -> None:
    text_before, _, text_after = placeholder_paragraph.text.partition(PLACEHOLDER)

    first_block, *remaining_blocks = blocks
    placeholder_paragraph.text = text_before if text_before.strip() else ""
    if not text_before.strip():
        placeholder_paragraph.style = _resolve_style_name(document, first_block.type)
        placeholder_paragraph.add_run(first_block.content)
        anchor = placeholder_paragraph
    else:
        anchor = _insert_paragraph_after(
            document,
            placeholder_paragraph,
            first_block.content,
            first_block.type,
        )

    for block in remaining_blocks:
        anchor = _insert_paragraph_after(document, anchor, block.content, block.type)

    if text_after.strip():
        _insert_paragraph_after(document, anchor, text_after, "paragraph")


def _replace_document_body(document: DocumentObject, blocks: list[StructuredBlock]) -> None:
    body = document._body._element
    children = list(body.iterchildren())
    content_start_index = _find_main_content_start_index(children)
    insertion_anchor_element = None

    for index, child in enumerate(children):
        if child.tag == qn("w:sectPr"):
            continue
        if content_start_index is None or index < content_start_index:
            insertion_anchor_element = child
            continue
        body.remove(child)

    if insertion_anchor_element is None:
        for block in blocks:
            paragraph = document.add_paragraph()
            paragraph.style = _resolve_style_name(document, block.type)
            paragraph.add_run(block.content)
        return

    for block in blocks:
        insertion_anchor_element = _insert_paragraph_after_element(
            insertion_anchor_element,
            document,
            block.content,
            block.type,
        )


def _find_main_content_start_index(children: list) -> int | None:
    for index, child in enumerate(children):
        if child.tag == qn("w:sectPr"):
            continue
        if _is_front_matter_element(child):
            continue
        return index
    return None


def _is_front_matter_element(element) -> bool:
    tag = element.tag
    if tag == qn("w:sdt"):
        return True
    if tag == qn("w:p"):
        text = "".join(node.text or "" for node in element.iter() if node.tag == qn("w:t")).strip()
        return not text
    return False


def _enable_field_updates_on_open(document: DocumentObject) -> None:
    settings_element = document.settings.element
    update_fields = settings_element.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings_element.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def _resolve_style_name(document: DocumentObject, block_type: str) -> str:
    available_styles = {style.name for style in document.styles}
    for candidate in STYLE_CANDIDATES[block_type]:
        if candidate in available_styles:
            return candidate
    return "Normal"


def _insert_paragraph_after(
    document: DocumentObject,
    paragraph: Paragraph,
    text: str,
    block_type: str,
) -> Paragraph:
    new_paragraph = OxmlElement("w:p")
    paragraph._element.addnext(new_paragraph)
    inserted = Paragraph(new_paragraph, paragraph._parent)
    inserted.style = _resolve_style_name(document, block_type)
    inserted.add_run(text)
    return inserted


def _insert_paragraph_after_element(element, document: DocumentObject, text: str, block_type: str):
    new_paragraph = OxmlElement("w:p")
    element.addnext(new_paragraph)
    inserted = Paragraph(new_paragraph, document._body)
    inserted.style = _resolve_style_name(document, block_type)
    inserted.add_run(text)
    return new_paragraph
