from __future__ import annotations

import base64
from datetime import date
from io import BytesIO
from pathlib import Path
import re

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from sqlalchemy.orm import Session

from webapp.dilmaria.exceptions import AgentExecutionError
from webapp.dilmaria.history_service import PopHistoryService
from webapp.dilmaria.pop_content_generator import PopContentGeneratorService
from webapp.dilmaria.pop_schema import (
    POP,
    PopDraftPreview,
    GeneratedPopContent,
    GuidedPopRequest,
    PopGenerationResult,
    PopRequest,
    PopSection,
    PopSubsection,
)
from webapp.dilmaria.pop_structures import get_pop_structure

ASSETS_DIR = Path(__file__).resolve().parent / "assets"
LOGO_PATH = ASSETS_DIR / "Logo_Naturale-removebg-preview.png"
NATURALE_TEMPLATE_PATH = ASSETS_DIR / "pop_naturale_modelo.docx"
MONTH_NAMES_PT_BR = {
    1: "janeiro",
    2: "fevereiro",
    3: "março",
    4: "abril",
    5: "maio",
    6: "junho",
    7: "julho",
    8: "agosto",
    9: "setembro",
    10: "outubro",
    11: "novembro",
    12: "dezembro",
}
TITLE_COLOR = RGBColor(28, 61, 107)
SUBTITLE_COLOR = RGBColor(88, 88, 88)
TEXT_COLOR = RGBColor(38, 38, 38)


async def run_pop_generator_agent(
    db: Session,
    user_id: int,
    raw_payload: dict,
) -> PopGenerationResult:
    request = await _resolve_request(raw_payload)
    history_service = PopHistoryService()
    structure = get_pop_structure(request.structure_key)
    pop = _build_pop(
        request,
        history_service.reserve_revision(db, request.codigo),
        structure.key,
        structure.name,
    )
    document = _build_document(pop)
    history_service.log_execution(db, pop, user_id=user_id, payload_snapshot=raw_payload)

    output = BytesIO()
    document.save(output)
    db.commit()
    return PopGenerationResult(pop=pop, document_bytes=output.getvalue())


async def preview_pop_generator_agent(raw_payload: dict) -> PopDraftPreview:
    guided_request = GuidedPopRequest.model_validate(raw_payload)
    structure = get_pop_structure(guided_request.structure_key)
    generator = PopContentGeneratorService()
    generated_content, refined_answers = await generator.build_draft(guided_request)
    draft = _merge_guided_request(guided_request, generated_content)
    return PopDraftPreview(
        draft=draft,
        refined_answers=refined_answers,
        structure_key=structure.key,
        structure_name=structure.name,
        creation_mode=guided_request.creation_mode,
    )


async def _resolve_request(raw_payload: dict) -> PopRequest:
    is_guided_payload = any(
        key in raw_payload for key in ("answers", "raw_context", "creation_mode")
    )
    if not is_guided_payload:
        raw_payload.setdefault("structure_key", "operacional_padrao")
        return PopRequest.model_validate(raw_payload)

    guided_request = GuidedPopRequest.model_validate(raw_payload)
    generated_content = await PopContentGeneratorService().build_content(guided_request)
    return _merge_guided_request(guided_request, generated_content)


def _merge_guided_request(guided_request: GuidedPopRequest, generated_content: GeneratedPopContent) -> PopRequest:
    return PopRequest(
        titulo=guided_request.titulo,
        codigo=guided_request.codigo,
        data=guided_request.data,
        custom_logo_data_url=guided_request.custom_logo_data_url,
        objetivo=generated_content.objetivo,
        documentos_referencia=generated_content.documentos_referencia,
        local_aplicacao=generated_content.local_aplicacao,
        responsabilidade_execucao=generated_content.responsabilidade_execucao,
        definicoes_siglas=generated_content.definicoes_siglas,
        atividades=generated_content.atividades,
        criterios_avaliacao=generated_content.criterios_avaliacao,
        boas_praticas=generated_content.boas_praticas,
        erros_criticos=generated_content.erros_criticos,
        termo=guided_request.termo,
        structure_key=guided_request.structure_key,
    )


def _build_pop(request: PopRequest, revisao: str, structure_key: str, structure_name: str) -> POP:
    pop_date = request.data or date.today()
    term = request.termo.model_copy(update={"data": request.termo.data or pop_date})
    display_title = PopContentGeneratorService()._normalize_fragment(  # noqa: SLF001
        request.titulo,
        lowercase_first=False,
    )

    activity_subsections: list[PopSubsection] = []
    for subsection_index, subsection in enumerate(request.atividades, start=1):
        subsection_number = f"6.{subsection_index}"
        cleaned_title = re.sub(r"^\d+(?:\.\d+)*\s+", "", subsection.titulo).strip()
        items = []
        for item_index, item in enumerate(subsection.itens, start=1):
            items.append(item.model_copy(update={"numero": f"{subsection_number}.{item_index}"}))
        activity_subsections.append(
            subsection.model_copy(
                update={
                    "numero": subsection_number,
                    "titulo": cleaned_title or subsection.titulo,
                    "itens": items,
                }
            )
        )

    sections = [
        PopSection(numero="1", titulo="Objetivo", conteudo=[request.objetivo]),
        PopSection(
            numero="2",
            titulo="Documentos de Referência",
            conteudo=request.documentos_referencia or ["Sem documentos de referência adicionais."],
        ),
        PopSection(numero="3", titulo="Local de Aplicação", conteudo=[request.local_aplicacao]),
        PopSection(
            numero="4",
            titulo="Responsabilidade / Execução",
            conteudo=[request.responsabilidade_execucao],
        ),
        PopSection(
            numero="5",
            titulo="Definições e Siglas",
            conteudo=[
                f"{definition.termo}: {definition.descricao}"
                for definition in request.definicoes_siglas
            ]
            or ["Não há definições ou siglas adicionais para este POP."],
        ),
        PopSection(
            numero="6",
            titulo="Descrição das Atividades",
            subsecoes=activity_subsections,
        ),
        PopSection(
            numero="7",
            titulo="Regras e Controle",
            subsecoes=[
                PopSubsection(
                    numero="7.1",
                    titulo="Critérios de Avaliação",
                    itens=[
                        {"numero": f"7.1.{index}", "descricao": value}
                        for index, value in enumerate(request.criterios_avaliacao, start=1)
                    ],
                ),
                PopSubsection(
                    numero="7.2",
                    titulo="Boas Práticas e Conformidade",
                    itens=[
                        {"numero": f"7.2.{index}", "descricao": value}
                        for index, value in enumerate(request.boas_praticas, start=1)
                    ],
                ),
                PopSubsection(
                    numero="7.3",
                    titulo="Erros Críticos / Falhas Proibidas",
                    itens=[
                        {"numero": f"7.3.{index}", "descricao": value}
                        for index, value in enumerate(request.erros_criticos, start=1)
                    ],
                ),
            ],
        ),
        PopSection(numero="8", titulo="Termo de Responsabilidade"),
    ]

    custom_logo_data = None if structure_key == "pop_naturale" else _extract_custom_logo_data(
        request.custom_logo_data_url
    )

    return POP(
        titulo=display_title,
        codigo=request.codigo,
        revisao=revisao,
        data=pop_date,
        secoes=sections,
        termo=term,
        logo_path=str(LOGO_PATH) if structure_key == "pop_naturale" and LOGO_PATH.exists() else None,
        logo_data=custom_logo_data,
        structure_key=structure_key,
        structure_name=structure_name,
    )


def _build_document(pop: POP) -> DocumentObject:
    if pop.structure_key == "pop_naturale":
        return _build_naturale_document(pop)
    return _build_standard_document(pop)


def _build_standard_document(pop: POP) -> DocumentObject:
    document = Document()
    _configure_document(document)
    _add_header_block(document, pop)
    _add_summary(document)
    document.add_page_break()

    for section in pop.secoes:
        _add_section_heading(document, section.numero, section.titulo)
        if section.numero == "6":
            for subsection in section.subsecoes:
                _add_activity_subsection(document, subsection)
        elif section.numero == "7":
            for subsection in section.subsecoes:
                _add_rule_subsection(document, subsection)
        elif section.numero == "8":
            _add_responsibility_term(document, pop)
        else:
            _add_text_block_list(document, section.conteudo)

    _enable_field_updates_on_open(document)
    return document


def _build_naturale_document(pop: POP) -> DocumentObject:
    if not NATURALE_TEMPLATE_PATH.exists():
        raise AgentExecutionError(
            f"Template da estrutura pop_naturale nao encontrado em {NATURALE_TEMPLATE_PATH}"
        )

    document = Document(NATURALE_TEMPLATE_PATH)
    _configure_document(document)
    _update_naturale_header(document, pop)
    anchor = _find_naturale_term_anchor(document)
    _populate_naturale_term(document, pop, anchor)
    declaration_anchor = anchor.insert_paragraph_before()
    declaration_anchor.add_run().add_break(WD_BREAK.PAGE)
    _insert_naturale_intro_before(declaration_anchor, pop)
    _insert_standard_pop_body_before(declaration_anchor, pop)
    _enable_field_updates_on_open(document)
    return document


def _configure_document(document: DocumentObject) -> None:
    for section in document.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(11)


def _set_paragraph_style(paragraph, preferred: str, fallback: str = "Normal") -> bool:
    try:
        paragraph.style = preferred
        return True
    except KeyError:
        try:
            paragraph.style = fallback
            return True
        except KeyError:
            return False


def _set_outline_level(paragraph, level: int) -> None:
    paragraph_properties = paragraph._p.get_or_add_pPr()
    outline = paragraph_properties.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        paragraph_properties.append(outline)
    outline.set(qn("w:val"), str(level))


def _apply_run_format(
    run,
    *,
    size: int,
    bold: bool = False,
    color: RGBColor = TEXT_COLOR,
    all_caps: bool = False,
) -> None:
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.all_caps = all_caps


def _format_title_paragraph(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.05


def _format_subtitle_paragraph(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.line_spacing = 1.0


def _format_summary_heading_paragraph(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.0


def _format_section_paragraph(paragraph, level: int) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.08


def _format_body_paragraph(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.2


def _format_label_paragraph(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.0


def _format_list_paragraph(paragraph, *, bullet: bool = False) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.left_indent = Cm(0.7)
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.1
    if bullet:
        paragraph.paragraph_format.left_indent = Cm(0.9)


def _format_toc_paragraph(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.line_spacing = 1.08


def _extract_custom_logo_data(data_url: str | None) -> bytes | None:
    if not data_url:
        return None
    try:
        _, encoded = data_url.split(",", 1)
        return base64.b64decode(encoded, validate=True)
    except Exception as exc:  # noqa: BLE001
        raise AgentExecutionError("Nao foi possivel ler a logo personalizada enviada.") from exc


def _add_pop_logo(run, pop: POP) -> None:
    if pop.logo_data:
        run.add_picture(BytesIO(pop.logo_data), width=Cm(3.5))
        return
    if pop.logo_path:
        run.add_picture(pop.logo_path, width=Cm(3.5))


def _add_header_block(document: DocumentObject, pop: POP) -> None:
    table = document.add_table(rows=1, cols=2)
    table.autofit = True
    left_cell, right_cell = table.rows[0].cells

    if pop.logo_path or pop.logo_data:
        logo_paragraph = left_cell.paragraphs[0]
        logo_run = logo_paragraph.add_run()
        _add_pop_logo(logo_run, pop)

    metadata_lines = [
        ("Título do Procedimento", pop.titulo),
        ("Código do POP", pop.codigo),
        ("Revisão", pop.revisao),
        ("Data", pop.data.strftime("%d/%m/%Y")),
        ("Estrutura", pop.structure_name),
    ]
    for index, (label, value) in enumerate(metadata_lines):
        paragraph = right_cell.paragraphs[0] if index == 0 else right_cell.add_paragraph()
        label_run = paragraph.add_run(f"{label}: ")
        label_run.bold = True
        paragraph.add_run(value)

    title = document.add_paragraph()
    title.style = "Title" if "Title" in {style.name for style in document.styles} else "Heading 1"
    title.alignment = 1
    title.add_run("Procedimento Operacional Padrão (POP)")


def _add_summary(document: DocumentObject) -> None:
    heading = document.add_paragraph()
    styled = _set_paragraph_style(heading, "Heading 1")
    _format_summary_heading_paragraph(heading)
    _set_outline_level(heading, 0)
    run = heading.add_run("Sumário")
    if not styled:
        _apply_run_format(run, size=15, bold=True, color=TITLE_COLOR)
    paragraph = document.add_paragraph()
    _format_toc_paragraph(paragraph)
    _append_toc_field(paragraph)


def _add_section_heading(document: DocumentObject, number: str, title: str) -> None:
    paragraph = document.add_paragraph()
    styled = _set_paragraph_style(paragraph, "Heading 1")
    _format_section_paragraph(paragraph, 1)
    _set_outline_level(paragraph, 0)
    run = paragraph.add_run(f"{number}. {title}")
    if not styled:
        _apply_run_format(run, size=14, bold=True, color=TITLE_COLOR)


def _add_text_block_list(document: DocumentObject, blocks: list[str]) -> None:
    for block in blocks:
        paragraph = document.add_paragraph()
        _set_paragraph_style(paragraph, "Normal")
        _format_body_paragraph(paragraph)
        run = paragraph.add_run(block)
        _apply_run_format(run, size=11)


def _add_activity_subsection(document: DocumentObject, subsection: PopSubsection) -> None:
    heading = document.add_paragraph()
    styled = _set_paragraph_style(heading, "Heading 2", fallback="Heading 1")
    _format_section_paragraph(heading, 2)
    _set_outline_level(heading, 1)
    run = heading.add_run(f"{subsection.numero} {subsection.titulo}")
    if not styled:
        _apply_run_format(run, size=12, bold=True, color=SUBTITLE_COLOR)

    _add_labeled_list(document, "Materiais Necessários", subsection.materiais)
    _add_labeled_list(document, "Preparação", subsection.preparacao)
    _add_labeled_list(document, "Etapas Iniciais", subsection.etapas_iniciais)

    if subsection.itens:
        step_heading = document.add_paragraph()
        _set_paragraph_style(step_heading, "Normal")
        _format_label_paragraph(step_heading)
        run = step_heading.add_run("Passos Operacionais")
        _apply_run_format(run, size=11, bold=True, color=SUBTITLE_COLOR)

    for item in subsection.itens:
        paragraph = document.add_paragraph()
        _set_paragraph_style(paragraph, "List Number", fallback="List Paragraph")
        _format_list_paragraph(paragraph)
        run = paragraph.add_run(f"{item.numero} {item.descricao}")
        _apply_run_format(run, size=11)
        if item.observacao:
            observation = document.add_paragraph()
            _set_paragraph_style(observation, "Normal")
            _format_body_paragraph(observation)
            label_run = observation.add_run("Observação: ")
            _apply_run_format(label_run, size=10, bold=True, color=SUBTITLE_COLOR)
            observation_run = observation.add_run(item.observacao)
            _apply_run_format(observation_run, size=10)


def _add_rule_subsection(document: DocumentObject, subsection: PopSubsection) -> None:
    heading = document.add_paragraph()
    styled = _set_paragraph_style(heading, "Heading 2", fallback="Heading 1")
    _format_section_paragraph(heading, 2)
    _set_outline_level(heading, 1)
    run = heading.add_run(f"{subsection.numero} {subsection.titulo}")
    if not styled:
        _apply_run_format(run, size=12, bold=True, color=SUBTITLE_COLOR)
    for item in subsection.itens:
        paragraph = document.add_paragraph()
        styled = _set_paragraph_style(paragraph, "List Bullet", fallback="List Paragraph")
        _format_list_paragraph(paragraph, bullet=True)
        run = paragraph.add_run(item.descricao if styled else f"- {item.descricao}")
        _apply_run_format(run, size=11)


def _add_responsibility_term(document: DocumentObject, pop: POP) -> None:
    paragraph = document.add_paragraph()
    _set_paragraph_style(paragraph, "Normal")
    _format_label_paragraph(paragraph)
    run = paragraph.add_run(f"Responsável: {pop.termo.nome_responsavel}")
    _apply_run_format(run, size=11, bold=True, color=SUBTITLE_COLOR)

    declaration = document.add_paragraph()
    _set_paragraph_style(declaration, "Normal")
    _format_body_paragraph(declaration)
    run = declaration.add_run(pop.termo.declaracao)
    _apply_run_format(run, size=11)

    signature_table = document.add_table(rows=3, cols=2)
    signature_table.style = "Table Grid"
    rows = [
        ("Elaboração", pop.termo.elaborado_por),
        ("Aprovação", pop.termo.aprovado_por),
        ("Local e Data", f"{pop.termo.local} - {pop.termo.data.strftime('%d/%m/%Y')}"),
    ]
    for row, values in zip(signature_table.rows, rows):
        row.cells[0].text = values[0]
        row.cells[1].text = values[1]


def _add_labeled_list(document: DocumentObject, label: str, items: list[str]) -> None:
    if not items:
        return
    paragraph = document.add_paragraph()
    _set_paragraph_style(paragraph, "Normal")
    _format_label_paragraph(paragraph)
    run = paragraph.add_run(label)
    _apply_run_format(run, size=11, bold=True, color=SUBTITLE_COLOR)
    for item in items:
        bullet = document.add_paragraph()
        styled = _set_paragraph_style(bullet, "List Bullet", fallback="List Paragraph")
        _format_list_paragraph(bullet, bullet=True)
        run = bullet.add_run(item if styled else f"- {item}")
        _apply_run_format(run, size=11)


def _append_toc_field(paragraph) -> None:
    begin_run = OxmlElement("w:r")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run.append(begin)

    instruction_run = OxmlElement("w:r")
    instruction = OxmlElement("w:instrText")
    instruction.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instruction.text = 'TOC \\o "1-3" \\h \\z \\u'
    instruction_run.append(instruction)

    separator_run = OxmlElement("w:r")
    separator = OxmlElement("w:fldChar")
    separator.set(qn("w:fldCharType"), "separate")
    separator_run.append(separator)

    placeholder_run = OxmlElement("w:r")
    placeholder_text = OxmlElement("w:t")
    placeholder_text.text = "Atualize o sumário no Word."
    placeholder_run.append(placeholder_text)

    end_run = OxmlElement("w:r")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run.append(end)

    paragraph._p.append(begin_run)
    paragraph._p.append(instruction_run)
    paragraph._p.append(separator_run)
    paragraph._p.append(placeholder_run)
    paragraph._p.append(end_run)


def _enable_field_updates_on_open(document: DocumentObject) -> None:
    settings_element = document.settings.element
    update_fields = settings_element.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings_element.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def _update_naturale_header(document: DocumentObject, pop: POP) -> None:
    header = document.sections[0].header
    if not header.tables:
        raise AgentExecutionError("O template da Naturale nao possui a tabela de cabecalho esperada.")

    table = header.tables[0]
    expected_rows = (
        (0, 1, pop.titulo),
        (1, 1, pop.titulo),
        (2, 1, pop.titulo),
        (0, 2, pop.codigo),
        (1, 2, f"Rev: {pop.revisao.replace('Rev.', '').replace('Rev', '').strip() or pop.revisao}"),
        (2, 2, f"Data: {pop.data.strftime('%d/%m/%Y')}"),
    )
    for row_index, cell_index, value in expected_rows:
        paragraph = table.rows[row_index].cells[cell_index].paragraphs[0]
        paragraph.text = value


def _find_naturale_term_anchor(document: DocumentObject):
    for paragraph in document.paragraphs:
        if paragraph.text.strip().lower() == "declaração de responsabilidade":
            return paragraph
    raise AgentExecutionError(
        "O template da Naturale nao possui o bloco de declaracao de responsabilidade esperado."
    )


def _populate_naturale_term(document: DocumentObject, pop: POP, anchor) -> None:
    paragraphs = document.paragraphs
    anchor_index = next(
        (index for index, paragraph in enumerate(paragraphs) if paragraph._p is anchor._p),
        None,
    )
    if anchor_index is None:
        raise AgentExecutionError(
            "Nao foi possivel localizar o bloco final de responsabilidade no template da Naturale."
        )
    declaration_paragraph = paragraphs[anchor_index + 1]
    declaration_text = declaration_paragraph.text
    declaration_text = re.sub(r"_{5,}", pop.termo.nome_responsavel, declaration_text, count=1)
    declaration_text = re.sub(r"_{5,}", pop.titulo, declaration_text, count=1)
    declaration_text = re.sub(r"Eu,\s*([^,]+?)\s+entendo", r"Eu, \1, entendo", declaration_text, count=1)
    declaration_paragraph.text = declaration_text

    date_paragraph = paragraphs[anchor_index + 7]
    date_paragraph.text = _format_naturale_date(pop)


def _format_naturale_date(pop: POP) -> str:
    month_name = MONTH_NAMES_PT_BR[pop.termo.data.month]
    return f"{pop.termo.local}, {pop.termo.data.day:02d} de {month_name} de {pop.termo.data.year}."


def _insert_naturale_intro_before(anchor, pop: POP) -> None:
    kicker = anchor.insert_paragraph_before()
    _format_subtitle_paragraph(kicker)
    kicker_run = kicker.add_run("PROCEDIMENTO OPERACIONAL PADRÃO")
    _apply_run_format(kicker_run, size=10, bold=True, color=SUBTITLE_COLOR, all_caps=True)

    title = anchor.insert_paragraph_before()
    _format_title_paragraph(title)
    title_run = title.add_run(pop.titulo)
    _apply_run_format(title_run, size=18, bold=True, color=TITLE_COLOR)

    support = anchor.insert_paragraph_before()
    _format_subtitle_paragraph(support)
    revision_label = pop.revisao.replace("Rev.", "Rev. ").replace("  ", " ").strip()
    support_run = support.add_run(
        f"Código {pop.codigo} | {revision_label} | Estrutura {pop.structure_name}"
    )
    _apply_run_format(support_run, size=10, color=SUBTITLE_COLOR)

    summary_heading = anchor.insert_paragraph_before()
    _format_summary_heading_paragraph(summary_heading)
    _set_outline_level(summary_heading, 0)
    summary_run = summary_heading.add_run("Sumário")
    _apply_run_format(summary_run, size=15, bold=True, color=TITLE_COLOR)

    toc_paragraph = anchor.insert_paragraph_before()
    _format_toc_paragraph(toc_paragraph)
    _append_toc_field(toc_paragraph)

    body_page_break = anchor.insert_paragraph_before()
    body_page_break.add_run().add_break(WD_BREAK.PAGE)


def _insert_standard_pop_body_before(anchor, pop: POP) -> None:
    sections_to_render = [section for section in pop.secoes if section.numero != "8"]
    for section in sections_to_render:
        _insert_section_heading_before(anchor, section.numero, section.titulo)
        if section.numero == "7":
            for subsection in section.subsecoes:
                _insert_rule_subsection_before(anchor, subsection)
        elif section.numero == "6":
            for subsection in section.subsecoes:
                _insert_activity_subsection_before(anchor, subsection)
        else:
            _insert_text_block_list_before(anchor, section.conteudo)


def _insert_section_heading_before(anchor, number: str, title: str):
    paragraph = anchor.insert_paragraph_before()
    styled = _set_paragraph_style(paragraph, "Heading 1")
    _format_section_paragraph(paragraph, 1)
    _set_outline_level(paragraph, 0)
    run = paragraph.add_run(f"{number}. {title}")
    if not styled:
        _apply_run_format(run, size=14, bold=True, color=TITLE_COLOR)
    return paragraph


def _insert_text_block_list_before(anchor, blocks: list[str]) -> None:
    for block in blocks:
        paragraph = anchor.insert_paragraph_before()
        _set_paragraph_style(paragraph, "Normal")
        _format_body_paragraph(paragraph)
        run = paragraph.add_run(block)
        _apply_run_format(run, size=11)


def _insert_activity_subsection_before(anchor, subsection: PopSubsection) -> None:
    heading = anchor.insert_paragraph_before()
    styled = _set_paragraph_style(heading, "Heading 2", fallback="Heading 1")
    _format_section_paragraph(heading, 2)
    _set_outline_level(heading, 1)
    run = heading.add_run(f"{subsection.numero} {subsection.titulo}")
    if not styled:
        _apply_run_format(run, size=12, bold=True, color=SUBTITLE_COLOR)

    _insert_labeled_list_before(anchor, "Materiais Necessários", subsection.materiais)
    _insert_labeled_list_before(anchor, "Preparação", subsection.preparacao)
    _insert_labeled_list_before(anchor, "Etapas Iniciais", subsection.etapas_iniciais)

    if subsection.itens:
        step_heading = anchor.insert_paragraph_before()
        _set_paragraph_style(step_heading, "Normal")
        _format_label_paragraph(step_heading)
        run = step_heading.add_run("Passos Operacionais")
        _apply_run_format(run, size=11, bold=True, color=SUBTITLE_COLOR)

    for item in subsection.itens:
        if item.observacao:
            observation = anchor.insert_paragraph_before()
            _set_paragraph_style(observation, "Normal")
            _format_body_paragraph(observation)
            label_run = observation.add_run("Observação: ")
            _apply_run_format(label_run, size=10, bold=True, color=SUBTITLE_COLOR)
            observation_run = observation.add_run(item.observacao)
            _apply_run_format(observation_run, size=10)
        paragraph = anchor.insert_paragraph_before()
        _set_paragraph_style(paragraph, "List Number", fallback="List Paragraph")
        _format_list_paragraph(paragraph)
        run = paragraph.add_run(f"{item.numero} {item.descricao}")
        _apply_run_format(run, size=11)


def _insert_rule_subsection_before(anchor, subsection: PopSubsection) -> None:
    heading = anchor.insert_paragraph_before()
    styled = _set_paragraph_style(heading, "Heading 2", fallback="Heading 1")
    _format_section_paragraph(heading, 2)
    _set_outline_level(heading, 1)
    run = heading.add_run(f"{subsection.numero} {subsection.titulo}")
    if not styled:
        _apply_run_format(run, size=12, bold=True, color=SUBTITLE_COLOR)

    for item in subsection.itens:
        paragraph = anchor.insert_paragraph_before()
        styled = _set_paragraph_style(paragraph, "List Bullet", fallback="List Paragraph")
        _format_list_paragraph(paragraph, bullet=True)
        run = paragraph.add_run(item.descricao if styled else f"- {item.descricao}")
        _apply_run_format(run, size=11)


def _insert_labeled_list_before(anchor, label: str, items: list[str]) -> None:
    if not items:
        return
    paragraph = anchor.insert_paragraph_before()
    _set_paragraph_style(paragraph, "Normal")
    _format_label_paragraph(paragraph)
    run = paragraph.add_run(label)
    _apply_run_format(run, size=11, bold=True, color=SUBTITLE_COLOR)
    for item in items:
        bullet = anchor.insert_paragraph_before()
        styled = _set_paragraph_style(bullet, "List Bullet", fallback="List Paragraph")
        _format_list_paragraph(bullet, bullet=True)
        run = bullet.add_run(item if styled else f"- {item}")
        _apply_run_format(run, size=11)
